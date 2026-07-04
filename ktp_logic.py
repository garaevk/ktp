# -*- coding: utf-8 -*-
"""
Core logic for KTP generator: extract lessons from docx,
build new docx with title page, generate dates.
"""

import os
import re
from datetime import datetime, date, timedelta
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

HAS_PDFPLUMBER = False
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    pass


# ===== DEFAULT DATE CONFIG FOR 2025-2026 =====

DEFAULT_QUARTERS = [
    {'name': 'I четверть',  'start': '2025-09-01', 'end': '2025-10-26'},
    {'name': 'II четверть', 'start': '2025-11-07', 'end': '2025-12-30'},
    {'name': 'III четверть', 'start': '2026-01-12', 'end': '2026-03-27'},
    {'name': 'IV четверть', 'start': '2026-04-06', 'end': '2026-05-26'},
]

DEFAULT_HOLIDAY_RANGES = [
    {'name': 'Осенние каникулы',  'start': '2025-10-27', 'end': '2025-11-06'},
    {'name': 'Зимние каникулы',   'start': '2025-12-31', 'end': '2026-01-11'},
    {'name': 'Весенние каникулы', 'start': '2026-03-28', 'end': '2026-04-05'},
]

DEFAULT_SPECIFIC_HOLIDAYS = [
    {'date': '2025-11-04', 'name': 'День народного единства'},
    {'date': '2025-11-06', 'name': 'Дополнительный выходной'},
    {'date': '2026-02-23', 'name': 'День защитника Отечества'},
    {'date': '2026-03-08', 'name': 'Международный женский день'},
    {'date': '2026-03-09', 'name': 'Выходной (перенос)'},
    {'date': '2026-05-01', 'name': 'Праздник Весны и Труда'},
    {'date': '2026-05-09', 'name': 'День Победы'},
]

WEEKDAYS_RU = ['Пн', 'Вт', 'Ср', 'Чт', 'Пт', 'Сб', 'Вс']
WEEKDAYS_FULL = ['Понедельник', 'Вторник', 'Среда', 'Четверг', 'Пятница', 'Суббота', 'Воскресенье']


# ===== TABLE DISCOVERY =====

def get_tables_info(filepath):
    """Return list of {index, rows, cols, header} for all tables."""
    doc = Document(filepath)
    info = []
    for i, table in enumerate(doc.tables):
        rows = len(table.rows)
        cols = len(table.rows[0].cells) if table.rows else 0
        first_cells = [cell.text.strip()[:40] for cell in table.rows[0].cells if cell.text.strip()]
        header = ' | '.join(first_cells[:4]) if first_cells else '(пустая шапка)'
        info.append({'index': i, 'rows': rows, 'cols': cols, 'header': header})
    return info


def auto_detect_table(filepath):
    """Try to find the table that looks like a lesson plan (has numbered rows)."""
    info = get_tables_info(filepath)
    best = None
    for t in info:
        if t['rows'] > 5 and t['cols'] >= 6:
            if best is None or t['rows'] > best['rows']:
                best = t
    return best['index'] if best else 0


def preview_rows(filepath, table_index, max_rows=8):
    """Return first rows of a table for preview (list of lists)."""
    doc = Document(filepath)
    if table_index >= len(doc.tables):
        return []
    table = doc.tables[table_index]
    result = []
    limit = min(len(table.rows), max_rows)
    for i in range(limit):
        cells = [cell.text.strip()[:60] for cell in table.rows[i].cells]
        result.append(cells)
    return result


# ===== LESSON EXTRACTION =====

def extract_lessons(filepath, table_index):
    """Extract lesson data from a table. Returns list of dicts."""
    doc = Document(filepath)
    if table_index >= len(doc.tables):
        return []

    table = doc.tables[table_index]
    lessons = []

    for i, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        if i < 2:
            continue
        num = cells[0]
        if not num or not num.lstrip('-').isdigit():
            continue
        if int(num) <= 0:
            continue
        lessons.append({
            'num': num,
            'theme': cells[1] if len(cells) > 1 else '',
            'hours': cells[2] if len(cells) > 2 else '',
            'control': cells[3] if len(cells) > 3 else '',
            'practical': cells[4] if len(cells) > 4 else '',
            'resources': cells[6] if len(cells) > 6 else '',
        })

    return lessons


# ===== DATE GENERATION =====

def parse_date_str(s):
    return datetime.strptime(s.strip(), '%Y-%m-%d').date()


def generate_dates(quarters, days_of_week, holiday_ranges, specific_holidays, total_lessons, lessons_per_day_map=None):
    """
    Generate date strings for lessons.

    Args:
        quarters: [{'start': '2025-09-01', 'end': '2025-10-26'}, ...]
        days_of_week: list of ints (0=Mon..6=Sun)
        holiday_ranges: [{'start': '2025-10-27', 'end': '2025-11-06'}, ...]
        specific_holidays: [{'date': '2025-11-04'}, ...]
        total_lessons: int
        lessons_per_day_map: dict {weekday: count} e.g. {1: 2, 3: 1} for Tue=2, Thu=1

    Returns:
        list of str (dd.mm.yyyy or empty string)
    """
    if lessons_per_day_map is None:
        lessons_per_day_map = {}

    holiday_set = set()
    for h in specific_holidays:
        try:
            holiday_set.add(parse_date_str(h['date']))
        except (ValueError, KeyError):
            pass
    for r in holiday_ranges:
        try:
            s = parse_date_str(r['start'])
            e = parse_date_str(r['end'])
            for n in range((e - s).days + 1):
                holiday_set.add(s + timedelta(n))
        except (ValueError, KeyError):
            pass

    all_dates = []
    for q in quarters:
        try:
            qs = parse_date_str(q['start'])
            qe = parse_date_str(q['end'])
            for n in range((qe - qs).days + 1):
                d = qs + timedelta(n)
                if d.weekday() in days_of_week and d not in holiday_set:
                    all_dates.append(d)
        except (ValueError, KeyError):
            pass

    result = []
    for date_obj in all_dates:
        count = lessons_per_day_map.get(date_obj.weekday(), 1)
        for _ in range(count):
            if len(result) >= total_lessons:
                break
            result.append(date_obj.strftime('%d.%m.%Y'))
        if len(result) >= total_lessons:
            break

    while len(result) < total_lessons:
        result.append('')
    return result


# ===== PDF CALENDAR PARSING =====

def parse_calendar_pdf(filepath):
    """
    Try to extract quarter dates and holidays from a PDF calendar schedule.
    Returns dict with quarters, holiday_ranges, specific_holidays or None on failure.
    """
    if not HAS_PDFPLUMBER:
        return None

    try:
        text_parts = []
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                txt = page.extract_text()
                if txt:
                    text_parts.append(txt)

        full_text = '\n'.join(text_parts)

        result = {
            'quarters': [],
            'holiday_ranges': [],
            'specific_holidays': [],
        }

        # Try to find quarter patterns
        quarter_patterns = [
            (r'I\s+четверть', r'II\s+четверть', r'III\s+четверть', r'IV\s+четверть'),
            (r'1\s+(?:ая\s+)?четверть', r'2\s+(?:ая\s+)?четверть',
             r'3\s+(?:ая\s+)?четверть', r'4\s+(?:ая\s+)?четверть'),
            (r'I\s+полугодие', r'II\s+полугодие'),
        ]

        # Try each pattern set
        for patterns in quarter_patterns:
            found = []
            for p in patterns:
                match = re.search(p, full_text, re.IGNORECASE)
                if match:
                    # Look for dates after this match
                    after = full_text[match.end():match.end() + 200]
                    dates_found = re.findall(r'(\d{1,2})[. ](\d{1,2})[. ](\d{4})', after)
                    if len(dates_found) >= 2:
                        d1 = f'{dates_found[0][2]}-{dates_found[0][1].zfill(2)}-{dates_found[0][0].zfill(2)}'
                        d2 = f'{dates_found[1][2]}-{dates_found[1][1].zfill(2)}-{dates_found[1][0].zfill(2)}'
                        found.append({'name': match.group().strip(), 'start': d1, 'end': d2})

            if found:
                result['quarters'] = found
                break

        # Fall back to defaults if no quarters found
        if not result['quarters']:
            return None

        # Holiday ranges
        holiday_names = ['каникул', 'выходн', 'праздн']
        for kw in holiday_names:
            for match in re.finditer(kw, full_text, re.IGNORECASE):
                after = full_text[match.start():match.start() + 150]
                dates_found = re.findall(r'(\d{1,2})[. ](\d{1,2})[. ](\d{4})', after)
                if len(dates_found) >= 2:
                    d1 = f'{dates_found[0][2]}-{dates_found[0][1].zfill(2)}-{dates_found[0][0].zfill(2)}'
                    d2 = f'{dates_found[1][2]}-{dates_found[1][1].zfill(2)}-{dates_found[1][0].zfill(2)}'
                    # Find a name before the match
                    before = full_text[max(0, match.start()-50):match.start()]
                    name_match = re.search(r'([А-Яа-я\s]+)', before[::-1])
                    name = name_match.group(0)[::-1].strip() if name_match else kw
                    result['holiday_ranges'].append({'name': name[:30], 'start': d1, 'end': d2})

        # Specific dates (standalone dates that might be holidays)
        all_dates = re.findall(r'(\d{1,2})[. ](\d{1,2})[. ](\d{4})', full_text)
        seen = set()
        for day, month, year in all_dates:
            ds = f'{year}-{month.zfill(2)}-{day.zfill(2)}'
            if ds not in seen:
                seen.add(ds)

        return result

    except Exception:
        return None


# ===== DOCUMENT BUILDING =====

def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=12):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'


def add_center(doc, text, size=14, bold=True, spacing=1.7):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = spacing
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = 'Times New Roman'
    return p


def build_ktp(lessons, dates, grade='', subject='', level='',
              control_count=8, practical_count=0, total_hours=None,
              school_name=''):
    """
    Create a .docx with title page and lesson table.

    Args:
        lessons: list of dicts from extract_lessons()
        dates: list of str (dd.mm.yyyy), one per lesson
        grade, subject, level: metadata for title page
        control_count, practical_count: for summary row
        total_hours: sum of all hours (auto if None)
        school_name: school name for title page (default: hardcoded)

    Returns:
        Document object
    """
    if not school_name:
        school_name = 'Полное наименование образовательной организации'

    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # ===== TITLE PAGE =====
    add_center(doc, 'МИНИСТЕРСТВО ПРОСВЕЩЕНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ', 14, True, 1.7)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(school_name)
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    add_center(doc, '\u200c\u200c\u200c ', 14, True, 1.7)
    add_center(doc, '\u041f\u041e\u0423\u0420\u041e\u0427\u041d\u041e\u0415 \u041f\u041b\u0410\u041d\u0418\u0420\u041e\u0412\u0410\u041d\u0418\u0415', 20, True, 1.7)
    add_center(doc, '', 14, True, 1.7)
    add_center(doc, '\u0420\u0410\u0411\u041e\u0427\u0410\u042f \u041f\u0420\u041e\u0413\u0420\u0410\u041c\u041c\u0410', 14, True, 1.7)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.7
    r = p.add_run('(ID 7784743)')
    r.bold = False
    r.font.size = Pt(14)
    r.font.name = 'Times New Roman'

    add_center(doc, '', 14, True, 1.7)

    subject_line = f'\u0443\u0447\u0435\u0431\u043d\u043e\u0433\u043e \u043f\u0440\u0435\u0434\u043c\u0435\u0442\u0430 \u00ab{subject}'
    if level:
        subject_line += f'. {level}'
    subject_line += '\u00bb'
    add_center(doc, subject_line, 14, True, 1.7)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.7
    r = p.add_run(f'\u0434\u043b\u044f \u043e\u0431\u0443\u0447\u0430\u044e\u0449\u0438\u0445\u0441\u044f {grade} \u043a\u043b\u0430\u0441\u0441\u043e\u0432')
    r.bold = False
    r.font.size = Pt(14)
    r.font.name = 'Times New Roman'

    for _ in range(4):
        add_center(doc, '', 14, True, 1.7)

    # ===== SECTION BREAK =====
    new_section = doc.add_section()
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = Cm(29.7)
    new_section.page_height = Cm(21.0)
    new_section.top_margin = Cm(3.0)
    new_section.bottom_margin = Cm(1.5)
    new_section.left_margin = Cm(2.0)
    new_section.right_margin = Cm(2.0)

    # ===== TABLE HEADER =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run('\u041a\u0430\u043b\u0435\u043d\u0434\u0430\u0440\u043d\u043e-\u0442\u0435\u043c\u0430\u0442\u0438\u0447\u0435\u0441\u043a\u043e\u0435 \u043f\u043b\u0430\u043d\u0438\u0440\u043e\u0432\u0430\u043d\u0438\u0435')
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f'\u043f\u043e {subject.lower()} \u0434\u043b\u044f {grade} \u043a\u043b\u0430\u0441\u0441\u0430')
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'

    # ===== TABLE =====
    num_rows = 2 + len(lessons) + 1
    table = doc.add_table(rows=num_rows, cols=7)
    table.style = 'Table Grid'

    col_widths = [Cm(0.9), Cm(10.5), Cm(1.5), Cm(2.0), Cm(2.0), Cm(2.0), Cm(5.5)]

    # Row 0 header
    h0 = [
        '\u2116 \u043f/\u043f', '\u0422\u0435\u043c\u0430 \u0443\u0440\u043e\u043a\u0430',
        '\u041a\u043e\u043b\u0438\u0447\u0435\u0441\u0442\u0432\u043e \u0447\u0430\u0441\u043e\u0432',
        '\u041a\u043e\u043b\u0438\u0447\u0435\u0441\u0442\u0432\u043e \u0447\u0430\u0441\u043e\u0432',
        '\u041a\u043e\u043b\u0438\u0447\u0435\u0441\u0442\u0432\u043e \u0447\u0430\u0441\u043e\u0432',
        '\u0414\u0430\u0442\u0430 \u043f\u0440\u043e\u0432\u0435\u0434\u0435\u043d\u0438\u044f',
        '\u042d\u043b\u0435\u043a\u0442\u0440\u043e\u043d\u043d\u044b\u0435 \u0446\u0438\u0444\u0440\u043e\u0432\u044b\u0435\n\u043e\u0431\u0440\u0430\u0437\u043e\u0432\u0430\u0442\u0435\u043b\u044c\u043d\u044b\u0435 \u0440\u0435\u0441\u0443\u0440\u0441\u044b',
    ]
    for i, text in enumerate(h0):
        set_cell(table.rows[0].cells[i], text, bold=True)

    # Row 1 header
    h1 = [
        '\u2116 \u043f/\u043f', '\u0422\u0435\u043c\u0430 \u0443\u0440\u043e\u043a\u0430',
        '\u0412\u0441\u0435\u0433\u043e',
        '\u041a\u043e\u043d\u0442\u0440\u043e\u043b\u044c\u043d\u044b\u0435\n\u0440\u0430\u0431\u043e\u0442\u044b',
        '\u041f\u0440\u0430\u043a\u0442\u0438\u0447\u0435\u0441\u043a\u0438\u0435\n\u0440\u0430\u0431\u043e\u0442\u044b',
        '\u0414\u0430\u0442\u0430 \u043f\u0440\u043e\u0432\u0435\u0434\u0435\u043d\u0438\u044f',
        '\u042d\u043b\u0435\u043a\u0442\u0440\u043e\u043d\u043d\u044b\u0435 \u0446\u0438\u0444\u0440\u043e\u0432\u044b\u0435\n\u043e\u0431\u0440\u0430\u0437\u043e\u0432\u0430\u0442\u0435\u043b\u044c\u043d\u044b\u0435 \u0440\u0435\u0441\u0443\u0440\u0441\u044b',
    ]
    for i, text in enumerate(h1):
        set_cell(table.rows[1].cells[i], text, bold=True)

    # Data rows
    for idx, lesson in enumerate(lessons):
        row = table.rows[idx + 2]
        dt = dates[idx] if idx < len(dates) else ''
        data = [
            lesson['num'], lesson['theme'], lesson['hours'],
            lesson['control'], lesson['practical'], dt, lesson['resources']
        ]
        for i, text in enumerate(data):
            align = WD_ALIGN_PARAGRAPH.LEFT if i == 1 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell(row.cells[i], text, align=align)

    # Summary row
    if total_hours is None:
        total_hours = sum(int(l['hours']) for l in lessons if l['hours'].isdigit())

    summary = ['\u0418\u0442\u043e\u0433\u043e', '',
               str(total_hours or len(lessons)),
               str(control_count), str(practical_count), '', '']
    for i, text in enumerate(summary):
        set_cell(table.rows[-1].cells[i], text, bold=True)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]

    return doc
